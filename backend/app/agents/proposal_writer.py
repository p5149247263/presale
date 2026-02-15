from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

try:
    from docx import Document
except Exception:
    Document = None  # type: ignore[assignment]

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.pdfgen import canvas
except Exception:
    A4 = None  # type: ignore[assignment]
    cm = None  # type: ignore[assignment]
    canvas = None  # type: ignore[assignment]

from app.models.schemas import (
    ArchitectureOption,
    CloudRecommendation,
    CostEstimate,
    ProposalDocument,
    ProposalSection,
    RequirementMatrix,
    RiskRegister,
)


class ProposalWriter:
    @staticmethod
    def _fallback_docx(path: Path, title: str, lines: list[str]) -> None:
        content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>"""
        rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""

        def p(txt: str) -> str:
            escaped = txt.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            return f"<w:p><w:r><w:t>{escaped}</w:t></w:r></w:p>"

        body = [p(title)] + [p(x[:250]) for x in lines]
        document = (
            "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
            "<w:document xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\">"
            f"<w:body>{''.join(body)}<w:sectPr/></w:body></w:document>"
        )
        with ZipFile(path, "w", ZIP_DEFLATED) as zf:
            zf.writestr("[Content_Types].xml", content_types)
            zf.writestr("_rels/.rels", rels)
            zf.writestr("word/document.xml", document)

    @staticmethod
    def _fallback_pdf(path: Path, lines: list[str]) -> None:
        body = ["BT", "/F1 10 Tf", "50 780 Td"]
        for idx, line in enumerate(lines[:70]):
            safe = line.replace("(", "[").replace(")", "]")
            if idx > 0:
                body.append("0 -14 Td")
            body.append(f"({safe[:110]}) Tj")
        body.append("ET")
        stream = "\n".join(body).encode("latin-1", errors="ignore")

        objects = []
        objects.append(b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n")
        objects.append(b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n")
        objects.append(
            b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >> endobj\n"
        )
        objects.append(b"4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n")
        objects.append(
            b"5 0 obj << /Length " + str(len(stream)).encode("ascii") + b" >> stream\n" + stream + b"\nendstream endobj\n"
        )

        content = bytearray(b"%PDF-1.4\n")
        offsets = [0]
        for obj in objects:
            offsets.append(len(content))
            content.extend(obj)

        xref_pos = len(content)
        content.extend(f"xref\n0 {len(objects)+1}\n".encode("ascii"))
        content.extend(b"0000000000 65535 f \n")
        for off in offsets[1:]:
            content.extend(f"{off:010d} 00000 n \n".encode("ascii"))
        content.extend(
            f"trailer << /Size {len(objects)+1} /Root 1 0 R >>\nstartxref\n{xref_pos}\n%%EOF\n".encode("ascii")
        )
        path.write_bytes(content)

    def build(
        self,
        project_id: str,
        project_name: str,
        reqs: RequirementMatrix,
        architectures: list[ArchitectureOption],
        cloud: CloudRecommendation,
        cost: CostEstimate,
        risks: RiskRegister,
    ) -> ProposalDocument:
        sections = [
            ProposalSection(
                title="Executive Summary",
                content="This proposal outlines a phased AI implementation with citation-grounded delivery and governance controls.",
                citations=reqs.requirements[0].citations if reqs.requirements else [],
            ),
            ProposalSection(
                title="Understanding of Problem",
                content="Client requires a secure, scalable AI platform with strict compliance and measurable business outcomes.",
                citations=reqs.requirements[0].citations if reqs.requirements else [],
            ),
            ProposalSection(
                title="Proposed Solution + Architecture",
                content=f"Primary approach: {architectures[0].name} with fallback option {architectures[1].name if len(architectures) > 1 else 'N/A'}.",
                citations=architectures[0].citations if architectures else [],
            ),
            ProposalSection(
                title="Implementation Plan",
                content="Phase 1 Discovery (2-4 weeks), Phase 2 Build (6-10 weeks), Phase 3 Scale/Operate (ongoing).",
                citations=cloud.citations,
            ),
            ProposalSection(
                title="Team and Delivery Model",
                content="Hybrid onshore/offshore team with PM, AI architect, ML engineers, data engineers, and QA.",
                citations=cloud.citations,
            ),
            ProposalSection(
                title="Pricing Assumptions and SOW Boundaries",
                content=(
                    f"Expected monthly run-rate: USD {cost.expected_total:,.2f}. Scope excludes source-system replatforming unless specified."
                ),
                citations=cost.citations,
            ),
            ProposalSection(
                title="Risks and Mitigations",
                content="Key risks include privacy, hallucination, and regulatory controls; mitigations include HITL and guardrails.",
                citations=risks.risks[0].citations if risks.risks else [],
            ),
            ProposalSection(
                title="Differentiators",
                content="Reusable accelerators, citation traceability, and measurable evaluation framework.",
                citations=cloud.citations,
            ),
        ]

        return ProposalDocument(
            project_id=project_id,
            title=f"Proposal - {project_name}",
            created_at=datetime.now(timezone.utc),
            sections=sections,
            appendix_requirements=reqs,
            appendix_risks=risks,
            appendix_architecture_mermaid=architectures[0].mermaid if architectures else "",
        )

    def export_docx(self, proposal: ProposalDocument, output_path: str) -> str:
        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        if Document is None:
            lines = [f"Created: {proposal.created_at.isoformat()}"]
            lines.extend([f"{s.title}: {s.content}" for s in proposal.sections])
            self._fallback_docx(path, proposal.title, lines)
            return str(path)

        doc = Document()
        doc.add_heading(proposal.title, level=0)
        doc.add_paragraph(f"Created: {proposal.created_at.isoformat()}")
        for section in proposal.sections:
            doc.add_heading(section.title, level=1)
            doc.add_paragraph(section.content)
            if section.citations:
                doc.add_paragraph(
                    "Citations: " + ", ".join([f"{c.doc_id} p{c.page} [{c.section}]" for c in section.citations])
                )

        doc.add_heading("Appendix - Requirements Matrix", level=1)
        for req in proposal.appendix_requirements.requirements:
            doc.add_paragraph(f"{req.requirement_id}: {req.text}")

        doc.add_heading("Appendix - Risk Register", level=1)
        for risk in proposal.appendix_risks.risks:
            doc.add_paragraph(f"{risk.risk_id} ({risk.category.value}): {risk.description}")

        doc.add_heading("Appendix - Architecture Mermaid", level=1)
        doc.add_paragraph(proposal.appendix_architecture_mermaid)
        doc.save(path)
        return str(path)

    def export_pdf(self, proposal: ProposalDocument, output_path: str) -> str:
        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        if canvas is None or A4 is None or cm is None:
            lines = [proposal.title, f"Created: {proposal.created_at.isoformat()}"]
            lines.extend([f"{s.title}: {s.content}" for s in proposal.sections])
            self._fallback_pdf(path, lines)
            return str(path)

        c = canvas.Canvas(str(path), pagesize=A4)
        width, height = A4
        y = height - 2 * cm

        c.setFont("Helvetica-Bold", 18)
        c.drawString(2 * cm, y, proposal.title)
        y -= 1 * cm
        c.setFont("Helvetica", 10)
        c.drawString(2 * cm, y, f"Created: {proposal.created_at.isoformat()}")
        y -= 1 * cm

        for section in proposal.sections:
            if y < 4 * cm:
                c.showPage()
                y = height - 2 * cm
            c.setFont("Helvetica-Bold", 13)
            c.drawString(2 * cm, y, section.title)
            y -= 0.7 * cm
            c.setFont("Helvetica", 10)
            text_obj = c.beginText(2 * cm, y)
            for line in (section.content[:600] + "\n").splitlines():
                text_obj.textLine(line[:105])
                y -= 0.45 * cm
            c.drawText(text_obj)

            if section.citations:
                c.setFont("Helvetica-Oblique", 9)
                citation_line = "Citations: " + ", ".join([f"{c2.doc_id} p{c2.page} [{c2.section}]" for c2 in section.citations])
                c.drawString(2 * cm, max(y, 2 * cm), citation_line[:120])
                y -= 0.6 * cm

        c.showPage()
        c.save()
        return str(path)
